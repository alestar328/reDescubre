# -*- coding: utf-8 -*-
"""Genera el Plan de Marketing de Re-descubre en .docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------- Paleta de marca ----------
PRIMARY = RGBColor(0x1F, 0x6F, 0x5C)   # verde-teal calmado
ACCENT  = RGBColor(0xE8, 0x8A, 0x3A)   # naranja cálido
DARK    = RGBColor(0x23, 0x2B, 0x2A)
GREY    = RGBColor(0x5F, 0x6B, 0x68)
LIGHT   = RGBColor(0xF2, 0xF6, 0xF4)

doc = Document()

# ---------- Estilos base ----------
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def set_cell_text(cell, text, bold=False, color=None, size=10, align='left', white=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {'left':WD_ALIGN_PARAGRAPH.LEFT,'center':WD_ALIGN_PARAGRAPH.CENTER,'right':WD_ALIGN_PARAGRAPH.RIGHT}[align]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if white:
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    elif color is not None:
        run.font.color.rgb = color

def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = PRIMARY
        # línea inferior
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6')
        bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'1F6F5C')
        pbdr.append(bottom); pPr.append(pbdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = ACCENT
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = DARK
    return p

def add_para(text, bold=False, italic=False, color=None, size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.color.rgb = PRIMARY
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_number(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
    p.add_run(text)
    return p

def make_table(headers, rows, widths=None, header_fill='1F6F5C'):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=10, white=True)
        shade(hdr[i], header_fill)
    for r_idx, row in enumerate(rows):
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
            if r_idx % 2 == 1:
                shade(cells[i], 'F2F6F4')
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t

# ============================================================
# PORTADA
# ============================================================
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('RE-DESCUBRE'); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = PRIMARY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Menos pantallas, más vida real'); r.italic = True; r.font.size = Pt(15); r.font.color.rgb = ACCENT

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PLAN DE MARKETING Y LANZAMIENTO'); r.bold = True; r.font.size = Pt(20); r.font.color.rgb = DARK
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Verano 2026 · Junio – Septiembre'); r.font.size = Pt(14); r.font.color.rgb = GREY

for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Actividades presenciales y saludables para jóvenes (12–25) · Barcelona')
r.font.size = Pt(12); r.font.color.rgb = GREY

for _ in range(6):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Documento de trabajo · editable'); r.font.size = Pt(10); r.font.color.rgb = GREY
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Versión 1.0 · 12 de junio de 2026'); r.font.size = Pt(10); r.font.color.rgb = GREY
doc.add_page_break()

# ============================================================
# 0. CÓMO USAR ESTE DOCUMENTO
# ============================================================
add_heading('Cómo usar este documento', 1)
add_para('Esta es una guía paso a paso pensada para una fase de validación: hasta septiembre estáis estudiando la '
         'viabilidad del proyecto, así que el plan está diseñado para conseguir el máximo alcance posible con un '
         'presupuesto bajo (50–300 €/mes) y, al mismo tiempo, recoger datos que confirmen o descarten que la idea '
         'funciona en Barcelona.')
add_para('Prioridad estratégica acordada: captar primero PROVEEDORES de actividades. Sin oferta no hay marketplace, '
         'así que el primer empuje va dirigido a quienes publican actividades, y solo después abrimos el grifo de la '
         'demanda (jóvenes y familias). El objetivo de cierre es un lanzamiento completo en septiembre, coincidiendo '
         'con la "vuelta al cole", el mejor momento del año para apuntarse a actividades.')
add_para('Cada fase incluye: objetivo, pasos concretos, herramientas digitales y métricas. Los importes son '
         'orientativos y editables. Las casillas "☐" son para que marquéis avances.')

# ============================================================
# 1. RESUMEN EJECUTIVO
# ============================================================
add_heading('1. Resumen ejecutivo', 1)
add_para('Re-descubre es un marketplace de actividades presenciales y saludables para jóvenes de 12 a 25 años en '
         'Barcelona, cuyo propósito es sacarlos de las pantallas y llevarlos a experiencias reales: deporte, '
         'naturaleza, arte, talleres y vida social sana.')
add_para('Durante el verano 2026 ejecutaremos un plan en 4 fases encadenadas:')
add_bullet('captar 15–30 proveedores fundadores y dejar el catálogo con oferta real.', 'Fase 1 (julio): ')
add_bullet('construir comunidad y lista de espera de jóvenes y familias en Barcelona.', 'Fase 2 (agosto): ')
add_bullet('lanzamiento público aprovechando la vuelta al cole, con campaña de pago.', 'Fase 3 (septiembre): ')
add_bullet('preparación y medición transversal durante todo el periodo.', 'Fase 0 + medición: ')
add_para('Con un presupuesto total estimado de 450–900 € para el verano, el plan combina contenido orgánico, '
         'colaboraciones locales y micro-campañas de pago muy segmentadas a Barcelona.')

# ============================================================
# 2. OBJETIVOS (SMART)
# ============================================================
add_heading('2. Objetivos del verano (SMART)', 1)
add_para('Objetivos medibles, alcanzables y con fecha. Ajustad las cifras a vuestra realidad: están calibradas para '
         'una marca nueva, local y con poco presupuesto.')
make_table(
    ['#', 'Objetivo', 'Métrica', 'Meta a 30/09', 'Tipo'],
    [
        ['1', 'Captar proveedores fundadores', 'Proveedores con actividad publicada', '15–30', 'Oferta'],
        ['2', 'Actividades reales en catálogo', 'Nº de actividades activas', '40–80', 'Oferta'],
        ['3', 'Lista de espera / registros', 'Emails o registros de usuarios', '300–800', 'Demanda'],
        ['4', 'Comunidad en redes', 'Seguidores Instagram + TikTok', '1.000–2.500', 'Marca'],
        ['5', 'Primeras reservas reales', 'Reservas completadas en la app', '50–150', 'Conversión'],
        ['6', 'Validación de hipótesis', 'Encuestas + entrevistas hechas', '10 prov. + 30 usuarios', 'Aprendizaje'],
    ],
    widths=[0.4, 2.4, 2.2, 1.1, 0.9],
)
add_para('')
add_para('Nota de validación: el objetivo nº 6 es el más importante para esta fase. El alcance no sirve de nada si no '
         'aprendéis por qué la gente se registra (o no) y si los proveedores ven valor real.', italic=True, color=GREY, size=10)

# ============================================================
# 3. PÚBLICO OBJETIVO
# ============================================================
add_heading('3. Público objetivo', 1)
add_para('Re-descubre es un marketplace de dos lados. Este verano el orden de captación es: primero oferta, luego demanda.')

add_heading('3.1 Lado OFERTA — Proveedores (prioridad 1)', 2)
add_para('Quién publica actividades: clubes y centros deportivos, escuelas de surf/escalada/yoga, asociaciones '
         'juveniles, casals y centros cívicos, monitores y profesionales independientes, academias de arte/música, '
         'gimnasios y entidades de ocio educativo.')
add_bullet('Más reservas y visibilidad sin coste de publicidad propio; llenar plazas vacías; llegar a un público joven.', 'Qué les motiva: ')
add_bullet('Desconfianza ante plataformas nuevas, miedo a comisiones, poco tiempo para dar de alta actividades.', 'Sus frenos: ')
add_bullet('"Te traemos jóvenes de Barcelona que buscan justo lo que ofreces. Publicar es gratis durante el lanzamiento."', 'Mensaje clave: ')

add_heading('3.2 Lado DEMANDA — Jóvenes y familias (prioridad 2)', 2)
add_para('Dos sub-públicos según el tipo de cuenta de la app:')
add_bullet('autónomos, deciden ellos mismos. Hablan TikTok/Instagram. Buscan planes, quedar, probar cosas nuevas.', 'Jóvenes 16–25 (explorer): ')
add_bullet('madres/padres que buscan actividades sanas para sus hijos de 12–17. Hablan Instagram, Facebook, grupos de colegio.', 'Familias (beneficiary): ')

# ============================================================
# 4. PROPUESTA DE VALOR Y MENSAJES
# ============================================================
add_heading('4. Propuesta de valor y mensajes clave', 1)
add_para('Una sola idea fuerza, repetida en todos los canales:', bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"Menos pantallas, más vida real."'); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=PRIMARY
add_para('Mensajes según a quién hablamos:')
make_table(
    ['Público', 'Mensaje principal', 'Llamada a la acción'],
    [
        ['Proveedores', 'Llena tus plazas con jóvenes que buscan justo tu actividad. Publicar es gratis.', 'Publica tu actividad gratis'],
        ['Jóvenes', 'Descubre planes reales en Barcelona y desconéctate del móvil.', 'Únete a la lista / Descúbrelo'],
        ['Familias', 'Actividades sanas y seguras para tus hijos, en un solo sitio.', 'Apúntate a la lista de espera'],
    ],
    widths=[1.1, 3.6, 1.8],
)

# ============================================================
# 5. PLAN PASO A PASO POR FASES
# ============================================================
add_heading('5. Plan paso a paso (junio → septiembre)', 1)

# --- FASE 0 ---
add_heading('Fase 0 — Preparación y cimientos (15–30 de junio)', 2)
add_para('Objetivo: dejar lista la base digital y de medición antes de empezar a captar. Sin esto, todo lo demás no se puede medir.', bold=True, color=GREY, size=10)
add_number('Definir la identidad mínima: logo (ya lo tenéis), paleta, 1 frase de marca y tono de voz.', '☐ ')
add_number('Reservar nombres de usuario iguales en todas las redes: @redescubre (Instagram, TikTok, y opcional Facebook/LinkedIn).', '☐ ')
add_number('Crear la landing de captación con dos botones: "Soy proveedor" y "Soy usuario / lista de espera".', '☐ ')
add_number('Instalar la medición: Google Analytics 4 + Meta Pixel + (opcional) TikTok Pixel en la web.', '☐ ')
add_number('Configurar un formulario/lista de emails (newsletter) y un formulario de alta de proveedores.', '☐ ')
add_number('Preparar 1 kit de prensa simple: qué es Re-descubre, para quién, capturas y logo en una carpeta compartida.', '☐ ')
add_number('Crear un panel sencillo (hoja de cálculo) para registrar proveedores contactados y su estado.', '☐ ')

# --- FASE 1 ---
add_heading('Fase 1 — Captación de proveedores (julio)', 2)
add_para('Objetivo: 15–30 proveedores fundadores con al menos una actividad publicada. Es la fase más importante: sin oferta no hay producto.', bold=True, color=GREY, size=10)
add_number('Hacer una lista de 80–100 proveedores objetivo en Barcelona (Google Maps, Instagram local, directorios de centros cívicos).', '☐ ')
add_number('Lanzar el "Programa Fundadores": ventajas por entrar ahora (0 % comisión durante X meses, posición destacada, sello "Fundador").', '☐ ')
add_number('Contacto directo 1 a 1: email + DM de Instagram + visita presencial a los más cercanos. El trato humano convierte mucho más que un anuncio.', '☐ ')
add_number('Onboarding asistido: ofrecedles dar de alta vosotros sus primeras actividades para quitarles fricción.', '☐ ')
add_number('Pedir a cada proveedor que comparta en SUS redes que ya está en Re-descubre (les pasáis plantilla e imágenes).', '☐ ')
add_number('Entrevistar a 8–10 proveedores: ¿qué les convence?, ¿qué comisión aceptarían?, ¿qué les falta? (validación).', '☐ ')
add_para('Herramientas: Gmail/Notion para seguimiento, Canva para el dossier de proveedor, WhatsApp Business para el trato directo.', italic=True, size=10, color=GREY)

# --- FASE 2 ---
add_heading('Fase 2 — Comunidad y demanda (agosto)', 2)
add_para('Objetivo: construir audiencia y lista de espera mientras se ultima el catálogo. Agosto es flojo en conversión pero ideal para crear contenido.', bold=True, color=GREY, size=10)
add_number('Activar Instagram + TikTok con publicación constante (ver calendario de contenidos, sección 7).', '☐ ')
add_number('Campaña "lista de espera" con incentivo: los primeros X registrados entran en un sorteo de una actividad gratis.', '☐ ')
add_number('Colaborar con 3–5 micro-influencers locales de Barcelona (ocio, deporte, vida sana) a cambio de actividad gratis o pago pequeño.', '☐ ')
add_number('Publicar contenido mostrando actividades reales ya cargadas: "10 planes para desconectar este verano en Barcelona".', '☐ ')
add_number('Entrar en grupos y comunidades locales: Reddit r/Barcelona, grupos de Facebook de familias, foros de ocio juvenil.', '☐ ')
add_number('Lanzar la primera micro-campaña de pago (20–40 €) para captar emails y testear qué mensaje engancha más.', '☐ ')
add_number('Encuestar a 30 usuarios potenciales: ¿qué actividades buscan?, ¿pagarían?, ¿qué les frena? (validación).', '☐ ')

# --- FASE 3 ---
add_heading('Fase 3 — Lanzamiento público (septiembre / vuelta al cole)', 2)
add_para('Objetivo: máximo alcance y primeras reservas reales, aprovechando que en septiembre todo el mundo busca actividades para el curso.', bold=True, color=GREY, size=10)
add_number('Fijar una fecha de "Día de lanzamiento" y avisar a toda la lista de espera 1 semana antes.', '☐ ')
add_number('Activar la campaña de pago principal del verano (el grueso del presupuesto) segmentada a Barcelona.', '☐ ')
add_number('Email + mensaje a la lista de espera con oferta de bienvenida y enlace directo a reservar.', '☐ ')
add_number('Nota de prensa a medios locales (Betevé, Time Out Barcelona, Ara, blogs de familias y de ocio juvenil).', '☐ ')
add_number('Pedir a los proveedores fundadores que empujen el lanzamiento el mismo día (efecto coordinado).', '☐ ')
add_number('Publicar testimonios y primeras reservas reales como prueba social.', '☐ ')
add_number('Cierre de verano: recopilar todos los datos y decidir la viabilidad con números reales (ver sección 9).', '☐ ')

# ============================================================
# 6. CANALES Y HERRAMIENTAS DIGITALES
# ============================================================
add_heading('6. Canales y herramientas digitales', 1)
add_para('Stack recomendado priorizando opciones gratuitas o muy baratas. Sustituid cualquiera por la que ya domináis.')
make_table(
    ['Necesidad', 'Herramienta recomendada', 'Coste', 'Para qué'],
    [
        ['Redes sociales', 'Instagram + TikTok', 'Gratis', 'Canal principal de marca y demanda'],
        ['Diseño de contenido', 'Canva (Pro opcional)', '0–12 €/mes', 'Posts, reels, dossier proveedores'],
        ['Programar publicaciones', 'Meta Business Suite / Metricool', '0–18 €/mes', 'Calendario y métricas de redes'],
        ['Email / newsletter', 'Mailchimp o Brevo (free tier)', '0–9 €/mes', 'Lista de espera y avisos'],
        ['Analítica web', 'Google Analytics 4', 'Gratis', 'Medir tráfico y conversiones'],
        ['Píxeles de anuncios', 'Meta Pixel + TikTok Pixel', 'Gratis', 'Medir y optimizar campañas'],
        ['Publicidad pagada', 'Meta Ads (IG/FB) + TikTok Ads', 'Presupuesto', 'Alcance y captación segmentada'],
        ['CRM / seguimiento', 'Notion o Google Sheets', 'Gratis', 'Proveedores contactados y estado'],
        ['Trato directo', 'WhatsApp Business', 'Gratis', 'Onboarding de proveedores'],
        ['Formularios / encuestas', 'Google Forms / Tally', 'Gratis', 'Validación y alta de proveedores'],
        ['SEO local', 'Google Business Profile', 'Gratis', 'Aparecer en búsquedas de Barcelona'],
        ['Acortar y medir enlaces', 'Bitly / enlaces UTM', 'Gratis', 'Saber qué canal trae registros'],
    ],
    widths=[1.4, 2.2, 1.0, 2.0],
)

# ============================================================
# 7. CALENDARIO DE CONTENIDOS
# ============================================================
add_heading('7. Calendario de contenidos (orgánico)', 1)
add_para('Ritmo sostenible para 1–2 personas: 3–4 publicaciones por semana. La constancia importa más que la cantidad. '
         'Reutilizad cada idea en formato reel (TikTok/IG) y post.')
make_table(
    ['Día', 'Tipo de contenido', 'Ejemplo'],
    [
        ['Lunes', 'Plan de la semana', '"3 planes para desconectar esta semana en Barcelona"'],
        ['Miércoles', 'Detrás de marca / propósito', 'Por qué creamos Re-descubre · datos de tiempo en pantallas'],
        ['Viernes', 'Actividad destacada / proveedor', 'Presentamos a un proveedor fundador y su actividad'],
        ['Domingo', 'Comunidad / interacción', 'Encuesta en stories: "¿Qué te gustaría probar?"'],
    ],
    widths=[1.0, 2.2, 3.4],
)
add_para('')
add_para('Pilares de contenido (repartid los temas entre ellos):', bold=True)
add_bullet('Inspiración: planes y actividades reales para hacer en Barcelona.')
add_bullet('Propósito: el problema de las pantallas y la salud mental juvenil (con datos).')
add_bullet('Producto: cómo funciona la app, lo fácil que es reservar / publicar.')
add_bullet('Comunidad: testimonios, proveedores, jóvenes, prueba social.')

# ============================================================
# 8. PRESUPUESTO
# ============================================================
add_heading('8. Presupuesto del verano (bajo)', 1)
add_para('Basado en 50–300 €/mes. Abajo, un reparto recomendado que concentra la inversión en septiembre, cuando hay '
         'más demanda y ya existe catálogo que ofrecer. Ajustad libremente.')
make_table(
    ['Concepto', 'Junio', 'Julio', 'Agosto', 'Septiembre', 'Total'],
    [
        ['Herramientas (Canva, email, etc.)', '15 €', '15 €', '15 €', '15 €', '60 €'],
        ['Publicidad pagada (Meta/TikTok)', '0 €', '30 €', '60 €', '180 €', '270 €'],
        ['Micro-influencers locales', '0 €', '0 €', '60 €', '60 €', '120 €'],
        ['Incentivos (sorteo / actividad gratis)', '0 €', '20 €', '40 €', '40 €', '100 €'],
        ['Imprevistos', '10 €', '15 €', '15 €', '20 €', '60 €'],
        ['TOTAL MENSUAL', '25 €', '80 €', '190 €', '315 €', '610 €'],
    ],
    widths=[2.6, 0.85, 0.85, 0.85, 1.0, 0.85],
)
add_para('')
add_para('Total estimado del verano: ~610 €. Si el presupuesto es más ajustado, recortad primero influencers e '
         'incentivos y mantened la publicidad de septiembre, que es la de mayor retorno.', italic=True, color=GREY, size=10)

# ============================================================
# 9. KPIs Y MEDICIÓN
# ============================================================
add_heading('9. KPIs y medición (revisión semanal)', 1)
add_para('Revisad estos números cada lunes. Lo que no se mide, no se puede mejorar ni validar.')
make_table(
    ['Indicador (KPI)', 'Qué mide', 'Dónde se ve'],
    [
        ['Proveedores activos', 'Salud de la oferta', 'Panel interno / Supabase'],
        ['Actividades publicadas', 'Profundidad del catálogo', 'Panel interno / Supabase'],
        ['Registros / lista de espera', 'Interés de la demanda', 'Mailchimp / Brevo'],
        ['Coste por registro (CPL)', 'Eficiencia de los anuncios', 'Meta Ads / TikTok Ads'],
        ['Seguidores y alcance', 'Notoriedad de marca', 'Instagram / TikTok insights'],
        ['Reservas completadas', 'Conversión real', 'App / Supabase'],
        ['Tasa de activación proveedor', '% que publican tras registrarse', 'Panel interno'],
    ],
    widths=[1.9, 2.6, 2.2],
)
add_para('')
add_para('Decisión de viabilidad (fin de septiembre): el proyecto da señales de viabilidad si se cumplen, al menos, '
         'tres de estas cuatro condiciones:', bold=True)
add_bullet('≥ 15 proveedores publicando de forma autónoma.')
add_bullet('≥ 300 registros con un coste por registro asumible (< 1,5 € aprox.).')
add_bullet('≥ 50 reservas reales completadas.')
add_bullet('Feedback cualitativo positivo y repetido en entrevistas (proveedores y usuarios).')

# ============================================================
# 10. RIESGOS
# ============================================================
add_heading('10. Riesgos y cómo mitigarlos', 1)
make_table(
    ['Riesgo', 'Impacto', 'Mitigación'],
    [
        ['Pocos proveedores (sin oferta)', 'Alto', 'Onboarding asistido + Programa Fundadores; priorizar julio'],
        ['Agosto vacacional (poca actividad)', 'Medio', 'Usarlo para crear contenido y lista, no para convertir'],
        ['Presupuesto limitado', 'Medio', 'Concentrar el gasto en septiembre; apoyarse en orgánico'],
        ['Equipo pequeño / poco tiempo', 'Alto', 'Calendario sostenible (3-4 posts/sem); automatizar y programar'],
        ['Mensaje que no engancha', 'Medio', 'Testear 2-3 mensajes con micro-campañas antes de septiembre'],
        ['Datos sin recoger', 'Alto', 'Instalar analítica y píxeles en la Fase 0, antes de captar'],
    ],
    widths=[2.2, 0.9, 3.6],
)

# ============================================================
# 11. CHECKLIST DE ARRANQUE
# ============================================================
add_heading('11. Checklist de arranque (esta semana)', 1)
add_para('Lo mínimo para empezar ya, sin esperar a tenerlo todo perfecto:')
for item in [
    'Crear/reservar @redescubre en Instagram y TikTok.',
    'Montar la landing con los dos botones (proveedor / usuario).',
    'Instalar Google Analytics 4 y Meta Pixel.',
    'Crear el formulario de lista de espera y el de alta de proveedores.',
    'Hacer la hoja de cálculo de proveedores objetivo (empezar con 20).',
    'Diseñar en Canva el dossier de "Programa Fundadores".',
    'Definir la fecha tentativa del lanzamiento de septiembre.',
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('☐  '); r.font.size = Pt(12); r.font.color.rgb = PRIMARY
    p.add_run(item)

# ============================================================
# 12. PRÓXIMOS PASOS
# ============================================================
add_heading('12. Próximos pasos sugeridos', 1)
add_number('Validar las cifras de objetivos y el presupuesto con vuestra realidad.', '')
add_number('Asignar responsable a cada fase (aunque seáis pocos, que cada bloque tenga dueño).', '')
add_number('Empezar por la checklist de arranque de la sección 11 esta misma semana.', '')
add_number('Volcar las tareas a una herramienta de seguimiento (Notion / Trello / Sheets).', '')
add_number('Reservar 30 min cada lunes para revisar KPIs y ajustar.', '')

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Re-descubre · Menos pantallas, más vida real'); r.italic=True; r.font.color.rgb=PRIMARY; r.font.size=Pt(11)

import os
out = r'D:\websites\redescubreTuMente\marketing\Plan_Marketing_Re-descubre_Verano2026.docx'
doc.save(out)
print('SAVED', out)
